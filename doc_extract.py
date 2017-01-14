import re
import os
import sys
import docx

def extract_IPII():
    num_docs = 67
    info_dir = 'data\\IPII'
    count = 0
    for filename in os.listdir(info_dir):
        assert filename.endswith(".docx")
        print count
        print filename
        extended_name = 'data\\IPII\\' + filename
        doc = docx.Document(extended_name)

        patient_num = re.findall(r'\d+', filename)
        print patient_num
        if(patient_num[0][2] == '1'):
            #contol
            f_sum = open('f_control.txt', 'w')
        else:
            assert(patient_num[0][2] == '2')
            f_sum = open('f_risk.txt', 'w')

        # create file for only patient responses
        text_name_patient = 'data\\IPII_patient\\' + filename.rstrip('.docx') + '_patient.txt'
        f_patient = open(text_name_patient, 'w')

        # create file for pairs of data
        text_name_pair = 'data\\IPII_pairs\\' + filename.rstrip('.docx') + '_pair.txt'
        f_pair = open(text_name_pair, 'w')

        prev = None

        for para in doc.paragraphs:
            if para.text:
                lower_str = para.text.lower().encode('ascii', 'ignore')
                cleaned_str = lower_str.lstrip('male voice: ')
                cleaned_str = cleaned_str.lstrip('female voice: ')
                cleaned_str = cleaned_str.lstrip('s: ')

                # check for new patient response
                if lower_str.startswith('male voice:') or lower_str.startswith('female voice:') or lower_str.startswith(
                        's:'):

                    f_pair.write('\n')
                    f_pair.write('subject: ' + cleaned_str)
                    f_patient.write('\n')
                    f_patient.write(cleaned_str)
                    f_sum.write('\n')
                    f_sum.write(cleaned_str)
                    prev = 'patient'

                # check for new interviewer response
                elif lower_str.startswith('interviewer:') or para.runs[0].italic == True:
                    f_pair.write('\n')
                    f_pair.write('i : ' + cleaned_str)
                    prev = 'interviewer'

                # check for continued answers in new paragraph
                else:
                    if prev == 'patient':
                        f_patient.write(cleaned_str)
                        f_pair.write('subject: ' + lower_str)
                        f_sum.write('\n')
                        f_sum.write(cleaned_str)

                    elif prev == 'interviewer':
                        f_pair.write('i: ' + lower_str)

        f_patient.close()
        f_pair.close()

        count += 1

    if count == num_docs:
        print 'Extraction Complete!'

def extract_EAR_Schizotypy():
    num_docs = 63
    info_dir = 'data\\EAR Schizotypy Study Transcripts'
    count = 0
    for filename in os.listdir(info_dir):
        assert filename.endswith(".docx")
        print count
        print filename
        extended_name = 'data\\EAR Schizotypy Study Transcripts\\' + filename
        doc = docx.Document(extended_name)
        #file name 2nd digit 1 = control 3 = Schizophrenia
        patient_num = re.findall(r'\d+', filename)
        print patient_num
        if(patient_num[1] == '14'):
            #3rd digit labels
            if(patient_num[0][2] == '1'):
                #contol
                f_sum = open('f_control.txt', 'a')
            else:
                assert(patient_num[0][2] == '2')
                f_sum = open('f_risk.txt', 'a')
        else:
            assert(patient_num[1] == '16')
            #2rd digit labels
            if(patient_num[0][1] == '1'):
                #contol
                f_sum = open('f_control.txt', 'a')
            else:
                assert(patient_num[0][1] == '2')
                f_sum = open('f_risk.txt', 'a')

        # create file for only patient responses
        text_name_patient = 'data\\EAR_Schizotypy_patient\\' + filename.rstrip('.docx') + '_patient.txt'
        f_patient = open(text_name_patient, 'w')

        # create file for pairs of data
        text_name_pair = 'data\\EAR_Schizotypy_pairs\\' + filename.rstrip('.docx') + '_pair.txt'
        f_pair = open(text_name_pair, 'w')

        prev = None

        for para in doc.paragraphs:
            if para.text:
                lower_str = para.text.lower().encode('ascii', 'ignore')
                line_arr = lower_str.split(':')
                if len(line_arr) > 1:
                    # check for new patient response
                    if lower_str.startswith('subject:') or lower_str.startswith('s:') or lower_str.startswith('s[') \
                            or lower_str.startswith('subject['):
                        f_pair.write('\n')
                        f_pair.write('subject: ' + line_arr[1])
                        f_patient.write('\n')
                        f_patient.write(line_arr[1])
                        f_sum.write('\n')
                        f_sum.write(line_arr[1])
                    # check for new interviewer response
                    elif not line_arr[0].startswith('recording'):
                        f_pair.write('\n')
                        f_pair.write('other voice: ' + line_arr[1])

        f_patient.close()
        f_pair.close()
        count += 1

    if count == num_docs:
        print 'Extraction Complete!'

def extract_EAR_Schizophrenia():
    num_docs = 18
    info_dir = 'data\\EAR Schizophrenia Pilot Study Transcripts'
    count = 0
    for filename in os.listdir(info_dir):
        assert filename.endswith(".docx")
        print count
        print filename
        #file name 2nd digit 1 = control 3 = Schizophrenia
        patient_num = re.findall(r'\d+', filename)
        print patient_num
        if(patient_num[0][1] == '1'):
            #contol
            f_sum = open('f_control.txt', 'a')
        else:
            assert(patient_num[0][1] == '3')
            f_sum = open('f_risk.txt', 'a')

        extended_name = 'data\\EAR Schizophrenia Pilot Study Transcripts\\' + filename
        doc = docx.Document(extended_name)

        # create file for only patient responses
        text_name_patient = 'data\\EAR_Schizophrenia_patient\\' + filename.rstrip('.docx') + '_patient.txt'
        f_patient = open(text_name_patient, 'w')

        # create file for pairs of data
        text_name_pair = 'data\\EAR_Schizophrenia_pairs\\' + filename.rstrip('.docx') + '_pair.txt'
        f_pair = open(text_name_pair, 'w')

        prev = None

        for para in doc.paragraphs:
            if para.text:
                lower_str = para.text.lower().encode('ascii', 'ignore')
                line_arr = lower_str.split(':')
                if len(line_arr) > 1:
                    # check for new patient response
                    if lower_str.startswith('subject:') or lower_str.startswith('s:') or lower_str.startswith('s[') \
                            or lower_str.startswith('subject['):
                        f_pair.write('\n')
                        f_pair.write('subject: ' + line_arr[1])
                        f_patient.write('\n')
                        f_patient.write(line_arr[1])
                        f_sum.write('\n')
                        f_sum.write(line_arr[1])

                    # check for new interviewer response
                    elif not line_arr[0].startswith('recording'):
                        f_pair.write('\n')
                        f_pair.write('other voice: ' + line_arr[1])

        f_patient.close()
        f_pair.close()
        count += 1

    if count == num_docs:
        print 'Extraction Complete!'

if __name__ == "__main__":
    sys.path.extend(['C:\\Users\\heyyj\\PycharmProjects\\Schizotypy'])
    #number of docx files in folder
    extract_IPII()
    extract_EAR_Schizotypy()
    extract_EAR_Schizophrenia()
